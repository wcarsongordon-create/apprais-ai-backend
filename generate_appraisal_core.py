#!/usr/bin/env python3
"""
AppraisAI Commercial Appraisal Report Generator
================================================
Generates a professional, USPAP-compliant commercial real estate appraisal
report as a formatted .docx file.

Usage:
  1. Modify the DATA SECTION below with property-specific information
  2. Run: python generate_appraisal.py
  3. Output: .docx file at OUTPUT_PATH

The script creates all sections, tables, formatting, and placeholder frames
automatically. Customize only the data variables.
"""

import os
import sys
from datetime import datetime

# Ensure python-docx is available
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    os.system("pip install python-docx --break-system-packages -q")
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    os.system("pip install Pillow --break-system-packages -q")
    from PIL import Image, ImageDraw, ImageFont



# ═══════════════════════════════════════════════════════════════════════════════
#  REPORT BUILDER — Variables are injected by generator.py at runtime
# ═══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1.0)
section.top_margin = section.bottom_margin = Inches(0.85)

# ── Color constants ──
NAVY = RGBColor(0x00, 0x35, 0xB0)
DARK = RGBColor(0x0F, 0x17, 0x2A)
TEAL = RGBColor(0x00, 0xAB, 0x80)
MUTED = RGBColor(0x6B, 0x72, 0x80)
ORANGE = RGBColor(0xB4, 0x55, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FILL_HEADER = 'D6E4FF'
FILL_ALT = 'F5F8FF'
FILL_ACCENT = 'EEF3FF'
FILL_DARK = '0F172A'
FILL_NAVY = '0035B0'
FILL_TEAL = '00AB80'


# ═══════════════════════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def shade_para(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def add_rule(color='0035B0', size='8'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.font.name = 'Arial'
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    add_rule()
    return p

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    return p

def para(text='', bold=False, italic=False, size=10, color=None, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = color
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def multi_run_para(runs_data, align=None, space_after=6):
    """Create paragraph with multiple differently-formatted runs.
    runs_data: list of dicts with keys: text, bold, italic, size, color"""
    p = doc.add_paragraph()
    for rd in runs_data:
        r = p.add_run(rd.get('text', ''))
        r.font.name = 'Arial'
        r.font.size = Pt(rd.get('size', 10))
        r.font.bold = rd.get('bold', False)
        r.font.italic = rd.get('italic', False)
        if rd.get('color'):
            r.font.color.rgb = rd['color']
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(space_after)
    return p

def placeholder(text):
    p = doc.add_paragraph()
    r = p.add_run(f'[{text}]')
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = ORANGE
    p.paragraph_format.space_after = Pt(4)
    return p

def field_row(label, value):
    p = doc.add_paragraph()
    r = p.add_run(f'{label}:  ')
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = DARK
    r2 = p.add_run(str(value))
    r2.font.name = 'Arial'
    r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(headers, rows, col_widths=None, header_fill=FILL_HEADER):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = DARK
        shade_cell(cell, header_fill)
        if col_widths:
            cell.width = Inches(col_widths[i])
    # Data
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = 'Arial'
            r.font.size = Pt(9)
            if r_idx % 2 == 1:
                shade_cell(cell, FILL_ALT)
            if col_widths:
                cell.width = Inches(col_widths[c_idx])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def photo_placeholder_img(filepath, label, width=640, height=420):
    """Create a modern placeholder image."""
    img = Image.new('RGB', (width, height), color=(238, 243, 255))
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, width-1, height-1], outline=(0, 53, 176), width=2)
    draw.line([(2, 2), (width-3, height-3)], fill=(200, 210, 230), width=1)
    draw.line([(width-3, 2), (2, height-3)], fill=(200, 210, 230), width=1)
    box_w, box_h = min(380, width-40), 80
    box_x = (width - box_w) // 2
    box_y = (height - box_h) // 2
    draw.rectangle([box_x, box_y, box_x+box_w, box_y+box_h], fill=(255,255,255), outline=(0, 53, 176), width=1)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 14)
        font_sm = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 11)
    except:
        font = ImageFont.load_default()
        font_sm = font
    bbox = draw.textbbox((0,0), label, font=font)
    tw = bbox[2] - bbox[0]
    draw.text(((width - tw) // 2, box_y + 20), label, fill=(0, 53, 176), font=font)
    sub = "Insert photo here"
    bbox2 = draw.textbbox((0,0), sub, font=font_sm)
    tw2 = bbox2[2] - bbox2[0]
    draw.text(((width - tw2) // 2, box_y + 48), sub, fill=(148, 163, 184), font=font_sm)
    img.save(filepath)
    return filepath

def add_photo_frame(label, caption, img_width=Inches(5.0)):
    """Add a photo placeholder frame to the document."""
    tmp_path = f'/tmp/_appraisai_ph_{label.replace(" ", "_").replace("/","_")[:30]}.png'
    photo_placeholder_img(tmp_path, label)
    p = para(f'Photo: {label}', bold=True, size=10, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(tmp_path, width=img_width)
    p.paragraph_format.space_after = Pt(2)
    p = para(caption, italic=True, size=8.5, color=MUTED, align='center', space_after=4)
    p2 = para(f'Date: {EFFECTIVE_DATE}  |  {PROPERTY_ADDRESS}, {MUNICIPALITY}, {STATE} {ZIP}',
              size=7.5, color=MUTED, align='center', space_after=16)
    return p2


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1: COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

# Top spacing
for _ in range(2):
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# AppraisAI brand
p = para('AppraisAI', align='center', size=28, bold=True, color=NAVY, space_after=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AI-Powered Commercial Real Estate Appraisals')
r.font.name = 'Arial'; r.font.size = Pt(11); r.font.color.rgb = TEAL; r.font.italic = True
p.paragraph_format.space_after = Pt(16)

add_rule('0035B0', '16')
para('', space_after=8)

# Report type
p = para(f'COMMERCIAL REAL ESTATE {REPORT_TYPE.upper()}', align='center', size=14, bold=True, color=DARK, space_after=4)
p = para(f'{PROPERTY_TYPE}', align='center', size=12, bold=False, color=NAVY, space_after=20)

# Property info block
cover_fields = [
    ('SUBJECT PROPERTY', f'{PROPERTY_ADDRESS}'),
    ('', f'{CITY}, {STATE} {ZIP}'),
    ('COUNTY / MUNICIPALITY', f'{COUNTY} County / {MUNICIPALITY}'),
    ('BLOCK / LOT', BLOCK_LOT),
    ('PROPERTY TYPE', PROPERTY_TYPE),
    ('', ''),
    ('OPINION OF VALUE', CONCLUDED_VALUE),
    ('EFFECTIVE DATE', EFFECTIVE_DATE),
    ('', ''),
    ('PREPARED FOR', CLIENT_NAME),
    ('', CLIENT_ORG),
    ('PREPARED BY', APPRAISER_NAME),
    ('', APPRAISER_FIRM),
    ('', APPRAISER_LICENSE),
    ('REPORT DATE', REPORT_DATE),
]

for label, val in cover_fields:
    if not label and not val:
        para('', space_after=8)
        continue
    if label:
        field_row(label, val)
    else:
        p = doc.add_paragraph()
        r = p.add_run(f'                            {val}')
        r.font.name = 'Arial'; r.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

# Footer note
para('', space_after=20)
p = para('CONFIDENTIAL', align='center', size=9, bold=True, color=MUTED, space_after=2)
p = para('This report is intended solely for the use of the client identified herein.', align='center', size=8, italic=True, color=MUTED)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2: LETTER OF TRANSMITTAL
# ═══════════════════════════════════════════════════════════════════════════════

h1('Letter of Transmittal')

para(REPORT_DATE, align='right')
para('')
para(f'{CLIENT_NAME}')
para(f'{CLIENT_ORG}')
para(f'{CLIENT_ADDRESS}')
para('')
para(f'Re: {REPORT_TYPE} — {PROPERTY_ADDRESS}, {CITY}, {STATE} {ZIP}')
para('')
para('Dear Client:')
para('')
para(f'At your request and authorization, we have prepared an appraisal of the {INTEREST_APPRAISED.lower()} interest in the above-referenced property. The accompanying research, analyses, and conclusions constitute an {REPORT_TYPE} compiled in accordance with Standard 2-2(a) of the Uniform Standards of Professional Appraisal Practice (USPAP).')
para('')
para(f'The purpose of this appraisal is to estimate the Market Value of the {INTEREST_APPRAISED.lower()} interest in the subject property, in "as is" condition, as of {EFFECTIVE_DATE}. The intended use of the report is for {INTENDED_USE}.')
para('')
para(f'The subject property is a {PROPERTY_TYPE.lower()} building located at {PROPERTY_ADDRESS}, {CITY}, {STATE} {ZIP}. The building contains approximately {GBA} square feet of gross building area on a {SITE_SF} SF ({SITE_ACRES} acre) site. The improvements were originally constructed {YEAR_BUILT} with renovations completed in {YEAR_RENOVATED}.')
para('')
para(f'Based upon our investigation and analysis, and subject to the assumptions and limiting conditions set forth in this report, it is our opinion that the Market Value of the subject property, as of {EFFECTIVE_DATE}, is:')
para('')

# Value callout
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(CONCLUDED_VALUE_WORDS)
r.font.name = 'Arial'; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'({CONCLUDED_VALUE})')
r.font.name = 'Arial'; r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(12)

para('This appraisal has been prepared in conformance with USPAP as promulgated by the Appraisal Standards Board of The Appraisal Foundation, and in compliance with Title XI of FIRREA. This appraisal was prepared with the assistance of AI technology through the AppraisAI platform; all conclusions and certifications represent the professional judgment of the signing appraiser.')
para('')
para('Respectfully submitted,')
para('')
placeholder(APPRAISER_NAME)
placeholder(APPRAISER_LICENSE)
para(APPRAISER_FIRM)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3: CERTIFICATION
# ═══════════════════════════════════════════════════════════════════════════════

h1('Appraiser Certification')

para('We certify that, to the best of our knowledge and belief:')
para('')

cert_items = [
    "The statements of fact contained in this report are true and correct.",
    "The reported analyses, opinions, and conclusions are limited only by the reported assumptions and limiting conditions and are our personal, impartial, and unbiased professional analyses, opinions, and conclusions.",
    "We have no present or prospective interest in the property that is the subject of this report and no personal interest with respect to the parties involved.",
    f"We have performed no prior services regarding the subject property within the three years preceding acceptance of this assignment, unless otherwise noted.",
    "We have no bias with respect to the property that is the subject of this report or to the parties involved with this assignment.",
    "Our engagement in this assignment was not contingent upon developing or reporting predetermined results.",
    "Our compensation for completing this assignment is not contingent upon the development or reporting of a predetermined value or direction in value that favors the cause of the client, the amount of the value opinion, the attainment of a stipulated result, or the occurrence of a subsequent event directly related to the intended use of this appraisal.",
    "Our analyses, opinions, and conclusions were developed, and this report has been prepared, in conformity with the Uniform Standards of Professional Appraisal Practice.",
    "The reported analyses, opinions, and conclusions were developed, and this report has been prepared, in conformity with the Code of Professional Ethics and Standards of Professional Appraisal Practice of the Appraisal Institute.",
    f"A personal inspection of the property that is the subject of this report was made on {INSPECTION_DATE}.",
    "No one provided significant real property appraisal assistance to the person(s) signing this certification, except as noted in the report.",
    "The use of this report is subject to the requirements of the Appraisal Institute relating to review by its duly authorized representatives.",
    "This appraisal was prepared with AI assistance through the AppraisAI platform. The signing appraiser has independently verified all data, analyses, and conclusions.",
    f"As of the date of this report, the signing appraiser has completed the continuing education requirements of the state licensing authority.",
]

for i, item in enumerate(cert_items):
    multi_run_para([
        {'text': f'{i+1}. ', 'bold': True, 'size': 9.5, 'color': NAVY},
        {'text': item, 'size': 9.5},
    ], space_after=4)

para('')
placeholder(f'{APPRAISER_NAME}')
placeholder(f'{APPRAISER_LICENSE}')
para(f'Date: {REPORT_DATE}', size=10)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4: EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

h1('Executive Summary')

summary_rows = [
    ['Location', f'{PROPERTY_ADDRESS}, {CITY}, {STATE} {ZIP}'],
    ['County / Municipality', f'{COUNTY} County / {MUNICIPALITY}'],
    ['Block / Lot', BLOCK_LOT],
    ['Census Tract', CENSUS_TRACT],
    ['Zoning / FAR', f'{ZONING} / {FAR}'],
    ['Property Type', PROPERTY_TYPE],
    ['Year Built / Renovated', f'{YEAR_BUILT} / {YEAR_RENOVATED}'],
    ['Stories', str(STORIES)],
    ['Gross Building Area', f'{GBA} SF'],
    ['Net Rentable Area', f'{NRA} SF'],
    ['Site Area', f'{SITE_SF} SF ({SITE_ACRES} AC)'],
    ['Parking', PARKING],
    ['Occupancy', OCCUPANCY],
    ['Condition', CONDITION],
    ['', ''],
    ['Flood Zone', FLOOD_ZONE],
    ['Interest Appraised', INTEREST_APPRAISED],
    ['', ''],
    ['VALUATION RECAP', ''],
    ['Sales Comparison Value', SALES_COMP_VALUE],
    ['Income Approach Value', INCOME_VALUE],
    ['Cost Approach Value', COST_VALUE],
    ['', ''],
    ['RECONCILED MARKET VALUE', CONCLUDED_VALUE],
    ['Value per SF (GBA)', VALUE_PER_SF],
    ['Overall Cap Rate', CAP_RATE],
    ['Net Operating Income', NOI],
    ['', ''],
    ['Effective Date', EFFECTIVE_DATE],
    ['Inspection Date', INSPECTION_DATE],
    ['Report Date', REPORT_DATE],
]

add_table(['Item', 'Detail'], summary_rows, col_widths=[2.8, 4.5])

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5: TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════

h1('Table of Contents')

toc_items = [
    'Letter of Transmittal',
    'Appraiser Certification',
    'Executive Summary',
    'Subject Property Photographs',
    'Area / Location Maps',
    'Purpose, Function & Dates',
    'Scope of Work',
    'Definition of Market Value',
    'Interest Appraised',
    'Regional / Area Analysis',
    'Neighborhood Analysis',
    'Site / Land Description',
    'Zoning Analysis',
    'Description of Improvements',
    'Real Estate Taxes',
    'Highest & Best Use',
    'Valuation Methodology',
    'Sales Comparison Approach',
    'Income Capitalization Approach',
    'Cost Approach',
    'Reconciliation & Final Value Opinion',
    'Assumptions & Limiting Conditions',
    'AI Technology Disclosure',
    'Addenda',
]

for item in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(item)
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6: SUBJECT PROPERTY PHOTOGRAPHS
# ═══════════════════════════════════════════════════════════════════════════════

h1('Subject Property Photographs')

para('The following photographs document the subject property. Placeholder frames are provided for insertion of actual inspection photographs by the reviewing appraiser.', italic=True, size=9, color=MUTED, space_after=12)

add_photo_frame('FRONT ELEVATION', f'Front elevation of {PROPERTY_ADDRESS} facing the primary street. Show full facade, signage, entry, and storefront.')
add_photo_frame('REAR ELEVATION', f'Rear elevation showing alley access, parking area ({PARKING}), and service entrance.')

doc.add_page_break()

add_photo_frame('STREET SCENE', f'Streetscape view showing subject in context with adjacent commercial uses and neighborhood character.')
add_photo_frame('INTERIOR - GROUND FLOOR', f'Ground-floor commercial space (~{NRA.split(",")[0] if "," in NRA else NRA} SF). Show layout, ceiling height, condition, and finishes.')

doc.add_page_break()

add_photo_frame('INTERIOR - UPPER FLOOR', f'Upper-floor space. Show layout, condition, and finishes.')
add_photo_frame('AERIAL / SATELLITE', f'Overhead view showing subject parcel ({SITE_SF} SF), building footprint, parking, and surrounding context.')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7: MAPS
# ═══════════════════════════════════════════════════════════════════════════════

h1('Area & Location Maps')

add_photo_frame('LOCATION MAP', f'Regional map showing subject at {PROPERTY_ADDRESS} relative to downtown, major highways, and landmarks.')
add_photo_frame('FLOOD MAP', f'FEMA FIRM Panel {FLOOD_PANEL} showing subject in {FLOOD_ZONE}.')

doc.add_page_break()

add_photo_frame('ZONING MAP', f'Zoning map showing subject in {ZONING} district with adjacent zone designations.')
add_photo_frame('COMPARABLE SALES MAP', f'Map showing subject and comparable sales 1-{len(COMPARABLE_SALES)} within the market area.')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 8: PURPOSE, FUNCTION & DATES
# ═══════════════════════════════════════════════════════════════════════════════

h1('Purpose, Function & Dates of Appraisal')

h2('Purpose')
para(f'The purpose of this appraisal is to estimate the Market Value of the {INTEREST_APPRAISED.lower()} interest in the subject property, in "as is" condition, as of the effective date of the appraisal.')

h2('Intended Use')
para(f'The intended use of this report is for {INTENDED_USE}. No other use is intended or authorized.')

h2('Intended User')
para(f'The intended user of this report is {CLIENT_NAME} ({CLIENT_ORG}). No other party may rely on this report without the express written consent of the appraiser.')

h2('Effective Date of Appraisal')
field_row('Effective Date of Value', EFFECTIVE_DATE)
field_row('Date of Inspection', INSPECTION_DATE)
field_row('Date of Report', REPORT_DATE)

h2('Property Rights Appraised')
para(f'{INTEREST_APPRAISED} — The {INTEREST_APPRAISED.lower()} interest represents the absolute ownership unencumbered by any other interest or estate, subject only to the limitations imposed by the governmental powers of taxation, eminent domain, police power, and escheat.')

h2('Report Format')
para(f'This report has been prepared as an {REPORT_TYPE} in accordance with Standard 2-2(a) of USPAP. This format provides a comprehensive presentation of the data, reasoning, and analyses used in developing the appraiser\'s opinion of value.')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 9: SCOPE OF WORK
# ═══════════════════════════════════════════════════════════════════════════════

h1('Scope of Work')

para('The scope of work for this appraisal assignment was determined by the intended use of the report, the needs of the client, the complexity of the property, and the availability of relevant data. The following steps were completed:')
para('')

scope_items = [
    f'Physical inspection of the subject property, including exterior and interior observations, conducted on {INSPECTION_DATE}.',
    'Inspection and analysis of the surrounding neighborhood and competitive market area.',
    f'Research and analysis of the {CITY} / {COUNTY} County real estate market, including review of comparable sales, rental data, and market trends.',
    'Application of the Sales Comparison Approach, Income Capitalization Approach, and Cost Approach to value.',
    'Reconciliation of the value indications derived from the applicable approaches into a final opinion of market value.',
    'Review of public records including deed transfers, tax assessments, zoning ordinances, and flood maps.',
    'AI-assisted data aggregation and analysis through the AppraisAI platform, with all outputs independently reviewed and verified by the signing appraiser.',
]

for item in scope_items:
    p = doc.add_paragraph()
    r = p.add_run(f'    {item}')
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

h2('Data Sources')
para(f'Data was gathered from public records, multiple listing services, commercial real estate databases, property tax records, the {COUNTY} County Assessor, market participants, and published economic data. All sources are identified within the relevant sections of this report.')

h2('Competency')
para(f'The appraiser is competent to perform this assignment based on education, experience, and familiarity with the {CITY} / {COUNTY} County commercial real estate market. The appraiser holds an active Certified General Real Estate Appraiser license and has completed all continuing education requirements.')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 10: DEFINITION OF MARKET VALUE
# ═══════════════════════════════════════════════════════════════════════════════

h1('Definition of Market Value')

para('The following definition of Market Value is used in this report, consistent with the definition provided by federal regulatory agencies pursuant to Title XI of FIRREA:')
para('')

# Indented definition block
def_text = (
    '"The most probable price which a property should bring in a competitive and open market under all '
    'conditions requisite to a fair sale, the buyer and seller each acting prudently and knowledgeably, '
    'and assuming the price is not affected by undue stimulus. Implicit in this definition is the '
    'consummation of a sale as of a specified date and the passing of title from seller to buyer under '
    'conditions whereby:'
)
p = para(def_text, italic=True, size=10, space_after=8)

conditions = [
    'Buyer and seller are typically motivated;',
    'Both parties are well informed or well advised, and acting in what they consider their own best interests;',
    'A reasonable time is allowed for exposure in the open market;',
    'Payment is made in terms of cash in U.S. dollars or in terms of financial arrangements comparable thereto; and',
    'The price represents the normal consideration for the property sold unaffected by special or creative financing or sales concessions granted by anyone associated with the sale."',
]

for cond in conditions:
    p = doc.add_paragraph()
    r = p.add_run(f'    {cond}')
    r.font.name = 'Arial'; r.font.size = Pt(10); r.font.italic = True
    p.paragraph_format.space_after = Pt(3)

para('')
para('Source: 12 CFR Part 34.42(g); 55 Federal Register 34696, August 24, 1990, as amended at 57 Federal Register 12202, April 9, 1992; 59 Federal Register 29499, June 7, 1994.', size=8.5, italic=True, color=MUTED)

h2('Exposure Time')
para('Exposure time is the estimated length of time that the property interest being appraised would have been offered on the market prior to the hypothetical consummation of a sale at market value on the effective date of the appraisal. Based on analysis of comparable sales and market conditions, the estimated exposure time for the subject property is 6 to 12 months.')

h2('Marketing Period')
para('The marketing period is a prospective estimate of the time required to sell the property at the concluded market value. Based on current market conditions, the estimated marketing period is 6 to 12 months.')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 11: INTEREST APPRAISED
# ═══════════════════════════════════════════════════════════════════════════════

h1('Interest Appraised')

para(f'The interest appraised in this report is the {INTEREST_APPRAISED} interest in the subject property.')
para('')

if INTEREST_APPRAISED.lower() == "fee simple":
    para('Fee Simple Estate is defined as: "Absolute ownership unencumbered by any other interest or estate, subject only to the limitations imposed by the governmental powers of taxation, eminent domain, police power, and escheat." (The Dictionary of Real Estate Appraisal, Appraisal Institute, 7th Edition)')
else:
    para('Leased Fee Interest is defined as: "The ownership interest held by the lessor, which includes the right to receive the contract rent specified in the lease plus the reversionary right when the lease expires." (The Dictionary of Real Estate Appraisal, Appraisal Institute, 7th Edition)')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 12: REGIONAL / AREA ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════

h1('Regional / Area Analysis')

for paragraph_text in REGIONAL_ANALYSIS.strip().split('\n\n'):
    para(paragraph_text.strip())

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 13: NEIGHBORHOOD ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════

h1('Neighborhood Analysis')

for paragraph_text in NEIGHBORHOOD_ANALYSIS.strip().split('\n\n'):
    para(paragraph_text.strip())

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 14: SITE / LAND DESCRIPTION
# ═══════════════════════════════════════════════════════════════════════════════

h1('Site / Land Description')

site_data = [
    ['Location', f'{PROPERTY_ADDRESS}, {CITY}, {STATE} {ZIP}'],
    ['Site Area', f'{SITE_SF} SF ({SITE_ACRES} AC)'],
    ['Shape', 'Rectangular (interior parcel)'],
    ['Topography', 'Level at street grade'],
    ['Street Frontage', f'Approx. 40 feet on Grant Avenue'],
    ['Access', f'Primary: {PROPERTY_ADDRESS.split()[1] if len(PROPERTY_ADDRESS.split()) > 1 else "Street"} Ave; Secondary: rear alley'],
    ['Utilities', 'Public water, sewer, gas, electric — all available and connected'],
    ['Flood Zone', f'{FLOOD_ZONE} (Panel: {FLOOD_PANEL})'],
    ['Environmental', 'No known environmental issues; Phase I not completed for this report'],
    ['Easements', 'Typical utility easements assumed; no adverse encroachments observed'],
]

add_table(['Characteristic', 'Description'], site_data, col_widths=[2.5, 4.8])

for paragraph_text in SITE_DESCRIPTION.strip().split('\n\n'):
    para(paragraph_text.strip())

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 15: ZONING ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════

h1('Zoning Analysis')

zoning_data = [
    ['Zoning District', ZONING],
    ['Maximum FAR', FAR],
    ['Permitted Uses', 'Retail, office, restaurant, personal services, mixed-use residential above ground floor'],
    ['Minimum Lot Size', 'No minimum specified for commercial'],
    ['Maximum Height', '35 feet / 2.5 stories'],
    ['Setbacks', 'Front: 0 ft (build-to line); Side: 0 ft; Rear: 10 ft'],
    ['Parking Required', '1 space per 400 SF commercial; 1 per dwelling unit'],
    ['Conformity', 'The subject property appears to be a legally conforming use under the current zoning designation.'],
]

add_table(['Zoning Item', 'Detail'], zoning_data, col_widths=[2.5, 4.8])

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 16: DESCRIPTION OF IMPROVEMENTS
# ═══════════════════════════════════════════════════════════════════════════════

h1('Description of Improvements')

improvement_data = [
    ['Property Type', PROPERTY_TYPE],
    ['Year Built', YEAR_BUILT],
    ['Year Renovated', YEAR_RENOVATED],
    ['Stories', str(STORIES)],
    ['Gross Building Area', f'{GBA} SF'],
    ['Net Rentable Area', f'{NRA} SF'],
    ['Construction', STRUCTURE],
    ['Foundation', FOUNDATION],
    ['Roof', ROOF],
    ['Exterior Walls', EXTERIOR_WALLS],
    ['HVAC', HVAC],
    ['Electrical', ELECTRICAL],
    ['Plumbing', PLUMBING],
    ['Interior Finishes', INTERIOR],
    ['Parking', PARKING],
    ['Overall Condition', CONDITION],
]

add_table(['Component', 'Description'], improvement_data, col_widths=[2.5, 4.8])

for paragraph_text in IMPROVEMENT_DESCRIPTION.strip().split('\n\n'):
    para(paragraph_text.strip())

h2('Effective Age & Remaining Economic Life')
para(f'The subject building was originally constructed {YEAR_BUILT} with substantial renovations in {YEAR_RENOVATED}. Given the quality of the recent renovation, the effective age is estimated at approximately 10 years. The estimated total economic life for this type of construction is 50 years, yielding a remaining economic life of approximately 40 years.')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 17: REAL ESTATE TAXES
# ═══════════════════════════════════════════════════════════════════════════════

h1('Real Estate Taxes')

tax_data = [
    ['Tax Jurisdiction', f'{MUNICIPALITY}, {COUNTY} County, {STATE}'],
    ['Block / Lot', BLOCK_LOT],
    ['Assessed Value — Land', '[Assessed Land Value]'],
    ['Assessed Value — Improvements', '[Assessed Improvement Value]'],
    ['Total Assessed Value', '[Total Assessed Value]'],
    ['Millage Rate', '[Millage Rate]'],
    ['Annual Tax', '[Annual Tax Amount]'],
    ['Tax Year', '[Current Tax Year]'],
    ['Tax Status', '[Current / Delinquent]'],
]

add_table(['Tax Item', 'Detail'], tax_data, col_widths=[2.8, 4.5])

para('Real estate tax data was obtained from the county assessor records. The appraiser has accepted the assessment at face value and has not performed an independent assessment analysis. If the assessment were to change significantly, it could affect the concluded value.', size=9.5)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 18: HIGHEST & BEST USE
# ═══════════════════════════════════════════════════════════════════════════════

h1('Highest & Best Use')

h2('As If Vacant')
for paragraph_text in HIGHEST_BEST_USE_VACANT.strip().split('\n\n'):
    para(paragraph_text.strip())

h2('As Improved')
for paragraph_text in HIGHEST_BEST_USE_IMPROVED.strip().split('\n\n'):
    para(paragraph_text.strip())

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 19: VALUATION METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════════

h1('Valuation Methodology')

para('Three generally accepted approaches may be used to estimate the market value of real property:')
para('')

h3('Sales Comparison Approach')
para('This approach compares the subject property to similar properties that have recently sold, making adjustments for differences in property characteristics. This approach is most applicable when sufficient comparable sales data is available and is a primary indicator of value for most property types.')

h3('Income Capitalization Approach')
para('This approach converts anticipated future income from the property into an indication of present value. This approach is particularly relevant for income-producing properties and is given significant weight when reliable income and expense data is available.')

h3('Cost Approach')
para('This approach estimates the cost to reproduce or replace the existing improvements, less depreciation, plus the value of the land. This approach is most applicable for newer properties or special-purpose properties where comparable sales and income data are limited.')

para('')
para(f'For this assignment, all three approaches have been developed and considered. Given the nature of the subject as an income-producing {PROPERTY_TYPE.lower()} property with adequate comparable sales data, primary emphasis is placed on the Sales Comparison and Income Capitalization approaches, with the Cost Approach providing additional support.')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 20: SALES COMPARISON APPROACH
# ═══════════════════════════════════════════════════════════════════════════════

h1('Sales Comparison Approach')

para('The Sales Comparison Approach develops an indication of value by comparing the subject property to similar properties that have recently sold in the competitive market. Adjustments are made for differences in property rights, financing, conditions of sale, market conditions (time), location, and physical characteristics.')

h2('Comparable Sales Summary')

comp_headers = ['Element', 'Subject'] + [f'Comp {c["num"]}' for c in COMPARABLE_SALES]
comp_rows = [
    ['Address', f'{PROPERTY_ADDRESS}'] + [c['address'].split(',')[0] for c in COMPARABLE_SALES],
    ['Sale Price', '—'] + [c['sale_price'] for c in COMPARABLE_SALES],
    ['Sale Date', '—'] + [c['sale_date'] for c in COMPARABLE_SALES],
    ['Property Type', PROPERTY_TYPE] + [c['property_type'] for c in COMPARABLE_SALES],
    ['GBA (SF)', GBA] + [c['gba'] for c in COMPARABLE_SALES],
    ['Site (SF)', SITE_SF] + [c['site_sf'] for c in COMPARABLE_SALES],
    ['Price/SF', '—'] + [c['price_per_sf'] for c in COMPARABLE_SALES],
    ['Year Built', YEAR_BUILT] + [c['year_built'] for c in COMPARABLE_SALES],
    ['Stories', str(STORIES)] + [str(c['stories']) for c in COMPARABLE_SALES],
    ['Condition', CONDITION] + [c['condition'] for c in COMPARABLE_SALES],
]

# Use smaller widths for comp grid
n_comps = len(COMPARABLE_SALES)
col_w = [1.8, 1.3] + [1.2] * n_comps
add_table(comp_headers, comp_rows, col_widths=col_w)

h2('Adjustment Grid')

adj_headers = ['Adjustment', 'Subject'] + [f'Comp {c["num"]}' for c in COMPARABLE_SALES]
adj_rows = [
    ['Sale Price', '—'] + [c['sale_price'] for c in COMPARABLE_SALES],
    ['Location', '—'] + [c['location_adj'] for c in COMPARABLE_SALES],
    ['Size', '—'] + [c['size_adj'] for c in COMPARABLE_SALES],
    ['Condition', '—'] + [c['condition_adj'] for c in COMPARABLE_SALES],
    ['Age/Quality', '—'] + [c['age_adj'] for c in COMPARABLE_SALES],
    ['Net Adjustment', '—'] + [c['net_adj'] for c in COMPARABLE_SALES],
    ['Adjusted Price', '—'] + [c['adjusted_price'] for c in COMPARABLE_SALES],
]

add_table(adj_headers, adj_rows, col_widths=col_w)

h2('Sales Comparison Conclusion')
para(f'The comparable sales, after adjustment, indicate a range of values for the subject property. The adjusted sale prices range from {min(c["adjusted_price"] for c in COMPARABLE_SALES)} to {max(c["adjusted_price"] for c in COMPARABLE_SALES)}. After considering the quality of the data, the degree of adjustment required, and the similarity of each comparable to the subject, the Sales Comparison Approach indicates a value for the subject of:')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Indicated Value by Sales Comparison: {SALES_COMP_VALUE}')
r.font.name = 'Arial'; r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(12)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 21: INCOME CAPITALIZATION APPROACH
# ═══════════════════════════════════════════════════════════════════════════════

h1('Income Capitalization Approach')

para('The Income Capitalization Approach converts anticipated future income into an indication of present value through Direct Capitalization. This method divides a stabilized Net Operating Income (NOI) by a market-derived capitalization rate.')

h2('Projected Operating Statement')

income_rows = [
    ['Potential Gross Income (PGI)', INCOME_DATA['pgi']],
    [f'  Less: Vacancy & Collection Loss ({INCOME_DATA["vacancy_rate"]})', f'({INCOME_DATA["vacancy_loss"]})'],
    ['Effective Gross Income (EGI)', INCOME_DATA['egi']],
    ['', ''],
    ['Operating Expenses:', ''],
    ['  Real Estate Taxes', INCOME_DATA['taxes']],
    ['  Insurance', INCOME_DATA['insurance']],
    ['  Maintenance & Repairs', INCOME_DATA['maintenance']],
    ['  Management (5% EGI)', INCOME_DATA['management']],
    ['  Reserves for Replacement', INCOME_DATA['reserves']],
    ['Total Operating Expenses', INCOME_DATA['total_expenses']],
    ['', ''],
    ['NET OPERATING INCOME (NOI)', INCOME_DATA['noi']],
]

add_table(['Line Item', 'Annual Amount'], income_rows, col_widths=[4.0, 2.5])

h2('Capitalization Rate Selection')
para(f'The overall capitalization rate was derived from analysis of comparable sales and market surveys. Published investor surveys for similar property types in comparable markets indicate cap rates ranging from 8.0% to 10.5%. Based on the quality, condition, and location of the subject property, a capitalization rate of {CAP_RATE} has been selected.')

h2('Direct Capitalization')

cap_rows = [
    ['Net Operating Income', INCOME_DATA['noi']],
    ['Capitalization Rate', CAP_RATE],
    ['Indicated Value (NOI / Cap Rate)', INCOME_VALUE],
]

add_table(['Component', 'Amount'], cap_rows, col_widths=[3.5, 3.0])

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Indicated Value by Income Approach: {INCOME_VALUE}')
r.font.name = 'Arial'; r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(12)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 22: COST APPROACH
# ═══════════════════════════════════════════════════════════════════════════════

h1('Cost Approach')

para('The Cost Approach estimates value by calculating the current cost to construct a replacement or reproduction of the existing improvements, deducting for all forms of depreciation, and adding the estimated land value.')

h2('Land Value')
para(f'Based on analysis of comparable land sales in the market area, the site value is estimated at [Land Value Estimate]. The site contains {SITE_SF} SF ({SITE_ACRES} acres) in the {ZONING} zone.')

h2('Cost Estimate')

cost_rows = [
    ['Estimated Land Value', '[Land Value]'],
    ['', ''],
    [f'Replacement Cost New ({GBA} SF x [$/SF])', '[RCN Amount]'],
    ['  Less: Physical Depreciation', '([Physical Depreciation])'],
    ['  Less: Functional Obsolescence', '([Functional Obsolescence])'],
    ['  Less: External Obsolescence', '([External Obsolescence])'],
    ['Depreciated Value of Improvements', '[Depreciated Improvement Value]'],
    ['', ''],
    ['INDICATED VALUE BY COST APPROACH', COST_VALUE],
]

add_table(['Component', 'Amount'], cost_rows, col_widths=[4.0, 2.5])

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Indicated Value by Cost Approach: {COST_VALUE}')
r.font.name = 'Arial'; r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(12)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 23: RECONCILIATION & FINAL VALUE
# ═══════════════════════════════════════════════════════════════════════════════

h1('Reconciliation & Final Value Opinion')

recon_rows = [
    ['Sales Comparison Approach', SALES_COMP_VALUE],
    ['Income Capitalization Approach', INCOME_VALUE],
    ['Cost Approach', COST_VALUE],
]

add_table(['Approach', 'Indicated Value'], recon_rows, col_widths=[4.0, 2.5])

para(f'The three approaches to value indicate a range from {SALES_COMP_VALUE} to {COST_VALUE}. In reconciling these indications, the following considerations apply:')
para('')
para(f'The Sales Comparison Approach is well supported by recent market activity in the {CITY} market area and reflects the behavior of buyers and sellers of similar properties. This approach is given significant weight.')
para('')
para(f'The Income Capitalization Approach provides a strong indication of value for this income-producing property and reflects investor expectations for similar properties. This approach is also given significant weight.')
para('')
para('The Cost Approach provides additional support but is given less emphasis due to the difficulty of accurately estimating depreciation for older, renovated properties.')
para('')

para(f'Based on the analysis contained in this report, and subject to the assumptions and limiting conditions herein, it is our opinion that the Market Value of the {INTEREST_APPRAISED.lower()} interest in the subject property, as of {EFFECTIVE_DATE}, is:')
para('')

# Final value callout - big and bold
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(CONCLUDED_VALUE_WORDS)
r.font.name = 'Arial'; r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'({CONCLUDED_VALUE})')
r.font.name = 'Arial'; r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(16)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 24: ASSUMPTIONS & LIMITING CONDITIONS
# ═══════════════════════════════════════════════════════════════════════════════

h1('Assumptions & Limiting Conditions')

for i, item in enumerate(ASSUMPTIONS_AND_CONDITIONS):
    multi_run_para([
        {'text': f'{i+1}. ', 'bold': True, 'size': 9.5, 'color': NAVY},
        {'text': item, 'size': 9.5},
    ], space_after=5)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 25: AI TECHNOLOGY DISCLOSURE
# ═══════════════════════════════════════════════════════════════════════════════

h1('AI Technology Disclosure')

# Accent box
p = para('', space_after=8)
shade_para(p, FILL_ACCENT)

para('This appraisal report was prepared with the assistance of artificial intelligence technology through the AppraisAI platform. In the interest of transparency and professional responsibility, the following disclosures are made:', bold=False, size=10, space_after=8)

ai_items = [
    ('AI-Assisted Data Aggregation', 'The AppraisAI platform was used to aggregate publicly available data including property records, comparable sales, market statistics, and economic indicators. All data has been independently reviewed and verified by the signing appraiser.'),
    ('AI-Assisted Analysis', 'Preliminary market analysis, comparable sale adjustments, and income projections were generated with AI assistance. The signing appraiser has independently evaluated all analyses and applied professional judgment to the final conclusions.'),
    ('Professional Responsibility', 'The use of AI technology does not diminish or transfer the professional responsibility of the signing appraiser. All opinions of value, certifications, and conclusions contained in this report represent the independent professional judgment of the signing Certified General Real Estate Appraiser.'),
    ('USPAP Compliance', 'This report has been prepared in full compliance with USPAP. The signing appraiser has ensured that all AI-generated content meets the standards of competency, ethical conduct, and professional practice required by USPAP and applicable state licensing regulations.'),
    ('Data Privacy', 'Client information and property data processed through the AppraisAI platform are handled in accordance with applicable data privacy laws and the confidentiality requirements of USPAP.'),
]

for title, text in ai_items:
    h3(title)
    para(text, space_after=6)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 26: ADDENDA
# ═══════════════════════════════════════════════════════════════════════════════

h1('Addenda')

para('The following items are included in the addenda to this report, as applicable:')
para('')

addenda_items = [
    'Engagement Letter / Scope of Work Agreement',
    'Comparable Sale Detail Sheets',
    'Comparable Rental Data',
    'Subject Property Legal Description',
    'Subject Property Deed',
    'Tax Assessment Records',
    'Flood Zone Determination',
    'Zoning Verification',
    'Appraiser Qualifications',
    'State Appraiser License',
]

for item in addenda_items:
    p = doc.add_paragraph()
    r = p.add_run(f'    {item}')
    r.font.name = 'Arial'; r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)

para('')
para('— End of Report —', align='center', bold=True, size=11, color=NAVY, space_after=20)

# AppraisAI footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Generated by AppraisAI')
r.font.name = 'Arial'; r.font.size = Pt(9); r.font.color.rgb = TEAL; r.font.italic = True
p.paragraph_format.space_after = Pt(2)
p = para('AI-Powered Commercial Real Estate Appraisals', align='center', size=8, italic=True, color=MUTED)


# ═══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════════════════════════

doc.save(OUTPUT_PATH)
print(f'Report saved to: {OUTPUT_PATH}')

# Cleanup temp placeholder images
import glob
for f in glob.glob('/tmp/_appraisai_ph_*.png'):
    try:
        os.remove(f)
    except:
        pass

print('Done.')
