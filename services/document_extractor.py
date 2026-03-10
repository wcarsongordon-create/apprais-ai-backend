"""
AppraisAI — Document Text Extractor
=====================================
Extracts readable text from client-uploaded documents (PDF, DOCX, XLSX, CSV, TXT)
so the content can be injected into the Claude research prompt.
"""

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# Max characters to extract per document (keeps prompt size manageable)
MAX_CHARS_PER_DOC = 12000


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from a document given its raw bytes and filename.
    Returns extracted text or a descriptive placeholder if extraction fails.
    """
    fname = filename.lower().strip()
    try:
        if fname.endswith(".pdf"):
            text = _extract_pdf(file_bytes)
        elif fname.endswith(".docx"):
            text = _extract_docx(file_bytes)
        elif fname.endswith((".xlsx", ".xls")):
            text = _extract_xlsx(file_bytes)
        elif fname.endswith((".txt", ".csv")):
            text = file_bytes.decode("utf-8", errors="replace")
        elif fname.endswith(".doc"):
            return f"[Legacy .doc format — {filename} — manual review required]"
        elif fname.endswith((".jpg", ".jpeg", ".png")):
            return f"[Image file — {filename} — manual review required]"
        else:
            return f"[Unsupported file type — {filename}]"

        # Truncate if very large
        if len(text) > MAX_CHARS_PER_DOC:
            text = text[:MAX_CHARS_PER_DOC] + f"\n\n[... truncated at {MAX_CHARS_PER_DOC} chars ...]"

        return text if text.strip() else f"[No extractable text found in {filename}]"

    except Exception as e:
        logger.warning(f"Text extraction failed for {filename}: {e}")
        return f"[Could not extract text from {filename}: {str(e)}]"


def _extract_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"[Page {i + 1}]\n{text.strip()}")
    return "\n\n".join(pages) if pages else "[PDF contained no extractable text]"


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines) if lines else "[DOCX contained no extractable text]"


def _extract_xlsx(file_bytes: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f"=== Sheet: {sheet.title} ===")
        row_count = 0
        for row in sheet.iter_rows(values_only=True):
            vals = [str(v) for v in row if v is not None and str(v).strip()]
            if vals:
                lines.append(" | ".join(vals))
                row_count += 1
                if row_count >= 500:  # Cap rows per sheet
                    lines.append("[... sheet truncated at 500 rows ...]")
                    break
    wb.close()
    return "\n".join(lines) if lines else "[XLSX contained no extractable data]"
